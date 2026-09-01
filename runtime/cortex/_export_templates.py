"""Export the live quotation templates and rate card to the terms folder on Drive.

Cortex is the LIVE source (quotation_presets, rate_card:<slug>); this folder is the human copy. Run
after any change to a preset, its terms, or a rate, so the two never drift apart. Idempotent: files
are updated in place, never duplicated.
"""
from __future__ import annotations

import sys

from docx import Document
from docx.shared import Pt, RGBColor

from . import documents, drive, quotation, ratecard, store

INK, MUT, TEAL = RGBColor(0x1A, 0x1A, 0x1A), RGBColor(0x5F, 0x6B, 0x70), RGBColor(0x0A, 0x7C, 0x8C)
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
TERMS_FOLDER = "159mEGnuBsWh_GqfPfPTf3q3zaNautfnX"


def _doc_for(key: str, preset: dict) -> str:
    d = Document()
    d.styles["Normal"].font.name = "Calibri"
    d.styles["Normal"].font.size = Pt(10.5)

    def run(t, size=10.5, bold=False, italic=False, color=None, font=None):
        r = d.add_paragraph().add_run(t)
        r.font.size, r.font.bold, r.font.italic = Pt(size), bold, italic
        r.font.color.rgb = color or INK
        if font:
            r.font.name = font

    run(f"SENSA QUOTATION TEMPLATE - {key.upper()}", 16, True, font="Poppins")
    run(f'Exported live from Cortex. Preset key: "{key}".', 9.5, italic=True, color=MUT)
    run("This is the reference copy of what Cortex prints on a quotation of this type. The LIVE version is "
        "the quotation_presets setting in Cortex: edit there, then re-export so the folder and the system "
        "never drift apart.", 9.5, italic=True, color=MUT)
    run("")
    run("Default title: " + str(preset.get("title") or ""), 10.5, bold=True)
    if preset.get("note"):
        run("Default note: " + preset["note"], 10)
    run("")
    run("DELIVERABLES BLOCK", 12, True, color=TEAL, font="Poppins")
    for x in preset.get("deliverables") or ["(none)"]:
        run("- " + str(x))
    run("")
    run("LINE-ITEM STRUCTURE", 12, True, color=TEAL, font="Poppins")
    for sec in preset.get("sections") or []:
        run(sec.get("header", ""), 10.5, bold=True)
        for it in sec.get("items", []):
            desc = it.get("desc") if isinstance(it, dict) else str(it)
            w = f"   (weight {it['weight']})" if isinstance(it, dict) and it.get("weight") else ""
            run("   - " + str(desc) + w)
    terms = preset.get("terms") or {}
    run("")
    run("TERMS PRINTED ON THE QUOTATION", 12, True, color=TEAL, font="Poppins")
    if terms.get("intro"):
        run(terms["intro"], 10, italic=True, color=MUT)
    for g in terms.get("groups") or []:
        run("")
        run(g.get("heading", ""), 11, True, color=TEAL)
        for line in g.get("lines") or []:
            run(line, 10)
    path = f"/tmp/template-{key}.docx"
    d.save(path)
    return path


def run(company: str = "sensa") -> str:
    out, tok = [], drive.access_token()
    presets = quotation.presets()
    for key, preset in presets.items():
        name = f"Sensa - Quotation Template - {key}.docx"
        data = open(_doc_for(key, preset), "rb").read()
        drive.upsert_in_folder(TERMS_FOLDER, name, DOCX, data, token=tok)
        for slug in ("sensa", "skyvision"):
            co = store.get_company_by_slug(slug)
            if co:
                documents.save(co["id"], slug, name, DOCX, data, uploaded_by="cortex export")
        out.append(name)
    card = ratecard.get(company)
    if card:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Rate card"
        ws.append(["Item", "Unit", "Budget (AED)", "Normal (AED)", "Notes"])
        for g in card.get("groups", []):
            ws.append([g.get("heading", "").upper(), "", "", "", ""])
            for it in g.get("items", []):
                ws.append([it.get("desc"), it.get("unit"), it.get("budget"), it.get("rate"), it.get("note")])
        for col, w in zip("ABCDE", (58, 16, 15, 15, 62)):
            ws.column_dimensions[col].width = w
        p = "/tmp/ratecard-export.xlsx"
        wb.save(p)
        nm = f"Sensa - Rate Card v{card.get('version', '1')}.xlsx"
        drive.upsert_in_folder(TERMS_FOLDER, nm, XLSX, open(p, "rb").read(), token=tok)
        out.append(nm)
    return "Exported to the terms folder on Drive: " + "; ".join(out)


if __name__ == "__main__":
    print(run(sys.argv[1] if len(sys.argv) > 1 else "sensa"))
